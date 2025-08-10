"""
Extrai conteúdo de um .docx em estrutura útil para IA.
- Headings (títulos e subtítulos)
- Parágrafos
- Listas (bullet/numeração)
- Tabelas
- Pares chave:valor
- Metadados (autor, criado, modificado)
- Texto bruto consolidado (raw_text)
"""

from docx import Document
from typing import Any, Dict, List
import re


def _heading_level(style_name: str | None) -> int | None:
    """
    Tenta descobrir o nível do heading a partir do nome do estilo.
    Ex.: "Heading 1", "Título 2" -> retorna 1 ou 2.
    Se não achar número, retorna None.
    """
    if not style_name:
        return None
    m = re.search(r"(\d+)$", style_name)
    return int(m.group(1)) if m else None


def extract_docx(path: str) -> Dict[str, Any]:
    """
    Lê um arquivo .docx e retorna um dicionário com os dados extraídos.
    """
    doc = Document(path)

    headings: List[str] = []
    paragraphs: List[str] = []
    key_values: List[Dict[str, str]] = []
    lists: List[List[str]] = []      # cada item é uma lista (bullet/numeração)
    tables: List[Dict[str, List[List[str]]]] = []
    outline_tree: List[Dict[str, Any]] = []  # hierarquia de seções (para referência)

    # Função auxiliar: detecta se o parágrafo é de lista (bullet/numeração)
    def is_list_para(p) -> bool:
        try:
            # Checa numeração no nível do XML (mais confiável)
            return p._p.pPr.numPr is not None  # type: ignore[attr-defined]
        except Exception:
            # Fallback por nome do estilo (menos confiável, mas ajuda)
            return p.style and ("List" in p.style.name or "Lista" in p.style.name)

    current_list: List[str] | None = None
    section_stack: List[Dict[str, Any]] = []  # pilha para montar a hierarquia de headings

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue

        style_name = p.style.name if p.style else ""

        # 1) Headings (títulos/subtítulos)
        if ("Heading" in style_name) or ("Título" in style_name):
            lvl = _heading_level(style_name) or 1
            headings.append(text)

            # Fecha lista corrente se estiver aberta
            if current_list:
                lists.append(current_list)
                current_list = None

            # Monta hierarquia simples (outline)
            while section_stack and section_stack[-1]["level"] >= lvl:
                section_stack.pop()

            node = {"level": lvl, "heading": text, "paragraphs": []}
            if section_stack:
                section_stack[-1].setdefault("children", []).append(node)
            else:
                outline_tree.append(node)
            section_stack.append(node)
            continue

        # 2) Listas
        if is_list_para(p):
            if current_list is None:
                current_list = []
            current_list.append(text)
        else:
            # Se estava em lista, fecha e salva
            if current_list:
                lists.append(current_list)
                current_list = None

            # 3) Parágrafos “comuns”
            paragraphs.append(text)

            # 4) Heurística simples para “Chave: Valor”
            if ":" in text:
                key, value = text.split(":", 1)
                if len(key) <= 60:  # evita falsos positivos muito longos
                    key_values.append({"key": key.strip(), "value": value.strip()})

        # Adiciona o parágrafo à última seção do outline (se existir)
        if section_stack:
            section_stack[-1]["paragraphs"].append(text)

    # Fecha lista no fim do documento (se necessário)
    if current_list:
        lists.append(current_list)

    # 5) Tabelas (cada célula como texto)
    for tb in doc.tables:
        rows = [[(c.text or "").strip() for c in r.cells] for r in tb.rows]
        tables.append({"rows": rows})

    # 6) Metadados do documento
    meta = {
        "author": doc.core_properties.author,
        "created": str(doc.core_properties.created),
        "modified": str(doc.core_properties.modified),
        "subject": doc.core_properties.subject,
        "category": doc.core_properties.category,
    }

    # 7) Título do documento
    title = doc.core_properties.title or (headings[0] if headings else None)

    # 8) Texto bruto consolidado (bom para prompt direto)
    raw_parts: List[str] = []
    if title:
        raw_parts.append(title)
    raw_parts.extend(headings)
    for lst in lists:
        raw_parts.extend(lst)
    raw_parts.extend(paragraphs)
    for tb in tables:
        raw_parts.extend(["\t".join(row) for row in tb["rows"]])
    raw_text = "\n".join([s for s in raw_parts if s])

    # 9) Outline “flattened” (lista) para facilitar consumo
    flat_outline: List[Dict[str, Any]] = []

    def walk(nodes: List[Dict[str, Any]]):
        for n in nodes:
            flat_outline.append(
                {
                    "level": n["level"],
                    "heading": n["heading"],
                    "paragraphs": n.get("paragraphs", []),
                }
            )
            walk(n.get("children", []))

    walk(outline_tree)

    return {
        "metadata": meta,
        "title": title,
        "raw_text": raw_text,
        "headings": headings,
        "key_values": key_values,
        "paragraphs": paragraphs,
        "lists": lists,
        "tables": tables,
        "outline": flat_outline,
    }
